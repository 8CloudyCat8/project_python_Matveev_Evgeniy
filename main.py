# MAKE_GRAPHS

from functools import cmp_to_key
import csv
import openpyxl
from openpyxl.styles import Font, Border, Side, Alignment
from openpyxl.styles.numbers import BUILTIN_FORMATS
import matplotlib.pyplot as plt
import numpy as np
import PIL.Image as Image
import os
from PIL import ImageGrab
import win32com.client
import docx
from docx.shared import Pt
from docx2pdf import convert


class Salary:
    def __init__(self, salaryFrom, salaryTo, salary_gross, salaryCurrency):
        self.salaryFrom = salaryFrom
        self.salaryTo = salaryTo
        self.salary_gross = salary_gross
        self.salaryCurrency = salaryCurrency

    def currency_to_rur(self):
        currency_to_rub = {
            "Тенге": 0.13,
            "Рубли": 1,
            "Гривны": 1.64,
            "Доллары": 60.66,
            "Узбекский сум": 0.0055,
            "Манаты": 35.68,
            "Белорусские рубли": 23.91,
            "Евро": 59.90,
            "Грузинский лари": 21.74,
            "Киргизский сом": 0.76,
        }
        return list(
            map(
                lambda m: int(m.replace(" ", ""))
                          * currency_to_rub[self.salaryCurrency],
                (self.salaryFrom, self.salaryTo),
            )
        )

    def get_salary(self):
        return sum(self.currency_to_rur()) / 2


class Vacancy:
    def __init__(
            self,
            name,
            description,
            skills,
            experience_id,
            premium,
            employer_name,
            salary,
            area_name,
            published_at,
    ):
        self.name = name
        self.description = description
        self.skills = skills
        self.experience_id = experience_id
        self.premium = premium
        self.employer_name = employer_name
        self.salary = salary
        self.area_name = area_name
        self.published_at = published_at


def csv_reader(file_name):
    with open(file_name, encoding="utf-8-sig") as p:
        reader = [k for k in csv.reader(p)]
        headers = reader.pop(0)
        inf = list(
            filter(lambda data: "" not in data and len(data) == len(headers), reader)
        )
    return headers, inf


def csv_filter(headers, inf):
    vacancies_full = []
    for row_inf in inf:
        list_inf = list(map(lambda m: row_inf[m], range(len(headers))))
        salary = Salary(list_inf[6], list_inf[7], list_inf[8], list_inf[9])
        skills = list_inf[2].split("__temp__")
        vacancy = Vacancy(
            list_inf[0],
            list_inf[1],
            skills,
            list_inf[3],
            list_inf[4],
            list_inf[5],
            salary,
            list_inf[10],
            list_inf[11],
        )
        vacancies_full.append(vacancy)
    return vacancies_full


def small_filter(inf):
    vacancies_full = []
    for row_inf in inf:
        salary = Salary(row_inf[1], row_inf[2], None, row_inf[3])
        vacancies_full.append(
            Vacancy(
                row_inf[0],
                None,
                None,
                None,
                None,
                None,
                salary,
                row_inf[4],
                row_inf[5],
            )
        )
    return vacancies_full


class DataSet:
    def __init__(self, file_name):
        (headers, inf) = csv_reader(file_name)
        vacancies_full = (
            small_filter(inf) if len(headers) <= 6 else csv_filter(headers, inf)
        )
        self.file_name = file_name
        self.vacancies_full_objects = vacancies_full


def formatter_info(vacancies_full):
    dic_cur = {
        "AZN": "Манаты",
        "BYR": "Белорусские рубли",
        "EUR": "Евро",
        "GEL": "Грузинский лари",
        "KGS": "Киргизский сом",
        "KZT": "Тенге",
        "RUR": "Рубли",
        "UAH": "Гривны",
        "USD": "Доллары",
        "UZS": "Узбекский сум",
    }

    def formatter_str_number(str_num):
        return str_num[: len(str_num) - 2] if -1 != str_num.find(".") else str_num

    def formatter_salary(value):
        salaryFrom = formatter_str_number(value.salaryFrom)
        salaryTo = formatter_str_number(value.salaryTo)
        salaryCurrency = dic_cur[value.salaryCurrency]
        return Salary(salaryFrom, salaryTo, None, salaryCurrency)

    def formatter_time(value):
        return value[0:4]

    for vacancy in vacancies_full:
        setattr(vacancy, "salary", formatter_salary(getattr(vacancy, "salary")))
        setattr(
            vacancy,
            "published_at",
            formatter_time(getattr(vacancy, "published_at")),
        )
    return vacancies_full


def finder_inf(vacancies_full, parameter):
    selected_v_y_count = {}
    city_level_salaries = {}
    count_vacancies_full_city = {}
    s_y_l = {}
    selected_s_y_l = {}
    count_vacancies_full_year = {}
    for vacancy in vacancies_full:
        salary = vacancy.salary.get_salary()
        year = int(vacancy.published_at)
        if year in s_y_l:
            sal_yr_lvl = s_y_l[year]
            s_y_l[year] = (sal_yr_lvl[0] + salary, sal_yr_lvl[1] + 1)
            count_vacancies_full_year[year] += 1
        else:
            s_y_l[year] = (salary, 1)
            count_vacancies_full_year[year] = 1
            selected_s_y_l[year] = (0, 0)
            selected_v_y_count[year] = 0
        if parameter in vacancy.name:
            sel_sal_ye_lvl = selected_s_y_l[year]
            selected_s_y_l[year] = (
                sel_sal_ye_lvl[0] + salary,
                sel_sal_ye_lvl[1] + 1,
            )
            selected_v_y_count[year] += 1
        if vacancy.area_name in city_level_salaries:
            sal_ct_lvl = city_level_salaries[vacancy.area_name]
            city_level_salaries[vacancy.area_name] = (
                sal_ct_lvl[0] + salary,
                sal_ct_lvl[1] + 1,
            )
            count_vacancies_full_city[vacancy.area_name] += 1
        else:
            count_vacancies_full_city[vacancy.area_name] = 1
            city_level_salaries[vacancy.area_name] = (salary, 1)
    return (
        s_y_l,
        selected_s_y_l,
        count_vacancies_full_year,
        selected_v_y_count,
        city_level_salaries,
        count_vacancies_full_city,
        len(vacancies_full),
    )


def print_inf(inf, req, t_or_g):
    def take_pairs(dictionary, n):
        count = 0
        list_res = []
        for j in range(len(dictionary)):
            list_res.append(dictionary.popitem())
            count += 1
            if count != n:
                continue
            break
        return dict(list_res)

    (
        s_y_l,
        selected_s_y_l,
        count_vacancies_full_year,
        selected_v_y_count,
        city_level_salaries,
        count_vacancies_full_city,
        vacancies_full_count,
    ) = inf
    (s_y_l, selected_s_y_l, city_level_salaries) = list(
        map(
            lambda dictionary: dict(
                map(
                    lambda dict_pair: (
                        dict_pair[0],
                        int(dict_pair[1][0])
                        if dict_pair[1][1] == 0
                        else int(dict_pair[1][0] / dict_pair[1][1]),
                    ),
                    dictionary.items(),
                )
            ),
            (s_y_l, selected_s_y_l, city_level_salaries),
        )
    )
    count_vacancies_full_city = dict(
        map(
            lambda dict_pair: (
                dict_pair[0],
                float(f"{dict_pair[1] / vacancies_full_count:.4f}"),
            ),
            count_vacancies_full_city.items(),
        )
    )

    def greater_than_or_equal(dict_pair):
        return dict_pair[1] >= 0.01

    count_vacancies_full_city = dict(
        filter(greater_than_or_equal, count_vacancies_full_city.items())
    )
    temp = [(key, value) for key, value in count_vacancies_full_city.items()]

    def function(x, y):
        return 1 if x[1] > y[1] else -1

    temp.sort(key=cmp_to_key(function))
    count_vacancies_full_city = dict(temp)
    city_level_salaries = dict(
        filter(
            lambda dict_pair: dict_pair[0] in count_vacancies_full_city,
            city_level_salaries.items(),
        )
    )
    temp = [(key, value) for key, value in city_level_salaries.items()]

    def function1(x, y):
        return 1 if x[1] > y[1] else -1

    temp.sort(key=cmp_to_key(function1))
    city_level_salaries = dict(temp)
    print("Динамика уровня зарплат по годам:", s_y_l)
    print("Динамика количества вакансий по годам:", count_vacancies_full_year)
    print(
        "Динамика уровня зарплат по годам для выбранной профессии:",
        selected_s_y_l,
    )
    print(
        "Динамика количества вакансий по годам для выбранной профессии:",
        selected_v_y_count,
    )
    pair_one = take_pairs(city_level_salaries, 10)
    print(
        "Уровень зарплат по городам (в порядке убывания):",
        pair_one,
    )
    pair_two = take_pairs(count_vacancies_full_city, 10)
    print(
        "Доля вакансий по городам (в порядке убывания):",
        pair_two,
    )

    def table():
        book = openpyxl.Workbook()
        book.remove(book.active)

        title1 = [
            "Год",
            "Средняя зарплата",
            f"Средняя зарплата - {req}",
            "Количество вакансий",
            f"Количество вакансий - {req}",
        ]

        title2 = ["Город", "Уровень зарплат", "", "Город", "Доля вакансий"]

        sheet_1 = book.create_sheet("Статистика по годам")
        sheet_2 = book.create_sheet("Статистика по городам")

        sheet_1.column_dimensions["A"].width = len(title1[0]) + 10
        sheet_1.column_dimensions["B"].width = len(title1[1]) + 30
        sheet_1.column_dimensions["C"].width = len(title1[2]) + 40
        sheet_1.column_dimensions["D"].width = len(title1[3]) + 30
        sheet_1.column_dimensions["E"].width = len(title1[4]) + 40

        sheet_2.column_dimensions["A"].width = len(title2[0]) + 32
        sheet_2.column_dimensions["B"].width = len(title2[1]) + 20
        sheet_2.column_dimensions["C"].width = len(title2[2]) + 4
        sheet_2.column_dimensions["D"].width = len(title2[3]) + 32
        sheet_2.column_dimensions["E"].width = len(title2[4]) + 20

        sheet_1.append(title1)
        sheet_2.append(title2)

        sheet_1["A1"].font = Font(bold=True, size=25)
        sheet_1["A1"].alignment = Alignment(horizontal='center')
        sheet_1["B1"].font = Font(bold=True, size=25)
        sheet_1["B1"].alignment = Alignment(horizontal='center')
        sheet_1["C1"].font = Font(bold=True, size=25)
        sheet_1["C1"].alignment = Alignment(horizontal='center')
        sheet_1["D1"].font = Font(bold=True, size=25)
        sheet_1["D1"].alignment = Alignment(horizontal='center')
        sheet_1["E1"].font = Font(bold=True, size=25)
        sheet_1["E1"].alignment = Alignment(horizontal='center')

        sheet_2["A1"].font = Font(bold=True, size=20)
        sheet_2["A1"].alignment = Alignment(horizontal='center')
        sheet_2["B1"].font = Font(bold=True, size=20)
        sheet_2["B1"].alignment = Alignment(horizontal='center')
        sheet_2["C1"].font = Font(bold=True, size=20)
        sheet_2["C1"].alignment = Alignment(horizontal='center')
        sheet_2["D1"].font = Font(bold=True, size=20)
        sheet_2["D1"].alignment = Alignment(horizontal='center')
        sheet_2["E1"].font = Font(bold=True, size=20)
        sheet_2["E1"].alignment = Alignment(horizontal='center')

        data1 = []
        for w in range(2007, 2022 + 1):
            data1.append(
                [
                    w,
                    s_y_l.get(w),
                    selected_s_y_l.get(w),
                    count_vacancies_full_year.get(w),
                    selected_v_y_count.get(w),
                ]
            )

        for i in data1:
            sheet_1.append(i)

        city_level_salaries_keys = list(pair_one.keys())

        city_level_salaries_values = list(pair_one.values())

        count_vacancies_full_city_keys = list(pair_two.keys())

        count_vacancies_full_city_values = list(pair_two.values())

        data2 = []
        for r in range(len(city_level_salaries_keys)):
            sector = [
                city_level_salaries_keys[r],
                city_level_salaries_values[r],
                "",
                count_vacancies_full_city_keys[r],
                count_vacancies_full_city_values[r],
            ]
            data2.append(sector)

        for i in data2:
            sheet_2.append(i)

        def set_border(ws, cell_range):
            rows = ws[cell_range]
            side = Side(border_style="medium", color="FF000000")

            rows = list(rows)
            for pos_y, cells in enumerate(rows):
                for pos_x, cell in enumerate(cells):
                    border = Border(
                        left=cell.border.left,
                        right=cell.border.right,
                        top=cell.border.top,
                        bottom=cell.border.bottom,
                    )
                    border.left = side
                    border.right = side
                    border.top = side
                    border.bottom = side

                    cell.border = border

        set_border(sheet_1, "A1:E17")
        set_border(sheet_2, "A1:B11")
        set_border(sheet_2, "D1:E11")

        for bor in range(2, 11 + 1):
            sheet_2[f"E{bor}"].number_format = BUILTIN_FORMATS[10]

        for i in range(2, 23):
            sheet_1[f"A{i}"].font = Font(size=25)
            sheet_1[f"B{i}"].font = Font(size=25)
            sheet_1[f"C{i}"].font = Font(size=25)
            sheet_1[f"D{i}"].font = Font(size=25)
            sheet_1[f"E{i}"].font = Font(size=25)

            sheet_2[f"A{i}"].font = Font(size=20)
            sheet_2[f"B{i}"].font = Font(size=20)
            sheet_2[f"C{i}"].font = Font(size=20)
            sheet_2[f"D{i}"].font = Font(size=20)
            sheet_2[f"E{i}"].font = Font(size=20)

        for p in range(2, 23):
            sheet_1[f"A{p}"].alignment = Alignment(horizontal='center')
            sheet_1[f"B{p}"].alignment = Alignment(horizontal='center')
            sheet_1[f"C{p}"].alignment = Alignment(horizontal='center')
            sheet_1[f"D{p}"].alignment = Alignment(horizontal='center')
            sheet_1[f"E{p}"].alignment = Alignment(horizontal='center')

            sheet_2[f"A{p}"].alignment = Alignment(horizontal='center')
            sheet_2[f"B{p}"].alignment = Alignment(horizontal='center')
            sheet_2[f"C{p}"].alignment = Alignment(horizontal='center')
            sheet_2[f"D{p}"].alignment = Alignment(horizontal='center')
            sheet_2[f"E{p}"].alignment = Alignment(horizontal='center')

        book.save("report.xlsx")
        book.close()
        if t_or_g == 'Вакансии':
            os.startfile(r'report.xlsx')
            exit(0)

    table()

    def diagram_one(a, b):
        req1 = req.lower()
        labels_int = list(a.keys())

        labels = [str(e) for e in labels_int]

        men_means = list(a.values())
        women_means = list(b.values())

        plt.rc("font", size=8)
        x = np.arange(len(labels))
        width = 0.35

        fig, ax = plt.subplots()
        ax.bar(x - width / 2, men_means, width, label="средняя з/п")
        ax.bar(x + width / 2, women_means, width, label=f"з/п {req1}")

        ax.set_title("Уровень зарплат по годам")
        ax.set_xticks(x, labels, rotation=90)
        ax.legend()
        ax.grid(axis="y")

        fig.tight_layout()

        plt.savefig("images/1.png", dpi=400)
        plt.show()

    diagram_one(s_y_l, selected_s_y_l)

    def diagram_two(a, b):
        req1 = req.lower()
        labels_int = list(a.keys())

        labels = [str(e) for e in labels_int]

        men_means = list(a.values())
        women_means = list(b.values())

        plt.rc("font", size=8)
        x = np.arange(len(labels))
        width = 0.35

        fig, ax = plt.subplots()
        ax.bar(x - width / 2, men_means, width, label="Количество вакансий")
        ax.bar(x + width / 2, women_means, width, label=f"Количество вакансий {req1}")

        ax.set_title("Количество вакансий по годам")
        ax.set_xticks(x, labels, rotation=90)
        ax.legend()
        ax.grid(axis="y")

        fig.tight_layout()

        plt.savefig("images/2.png", dpi=400)
        plt.show()

    diagram_two(count_vacancies_full_year, selected_v_y_count)

    def diagram_three(a):
        vals = list(a.keys())
        for n, i in enumerate(vals, 0):
            if " " in i:
                i = i.replace(" ", "\n")
                vals[n] = i
        vals.reverse()
        values = list(a.values())
        values.reverse()
        index = np.arange(10)
        fig, ax = plt.subplots()
        plt.rc("font", size=8)
        plt.title("Уровень зарплат по городам")
        plt.barh(index, values)
        plt.yticks(index, vals)
        ax.grid(axis="x")
        fig.tight_layout()
        plt.savefig("images/3.png", dpi=400)
        plt.show()

    diagram_three(pair_one)

    def diagram_four(a):
        labels = list(a.keys())
        labels.insert(0, "Другие")
        values = list(a.values())
        values.insert(0, 1 - sum(values))
        fig, ax = plt.subplots()
        plt.rc("font", size=6)
        plt.pie(values, labels=labels)
        plt.axis("equal")
        fig.tight_layout()
        plt.savefig("images/4.png", dpi=400)
        plt.show()

    diagram_four(pair_two)


def save_all_diagrams():
    IMAGES_PATH = "images/"
    IMAGES_FORMAT = [".png"]
    IMAGE_SIZE_1 = 6400
    IMAGE_SIZE_2 = 4800
    IMAGE_ROW = 2
    IMAGE_COLUMN = 2

    image_names = [
        name
        for name in os.listdir(IMAGES_PATH)
        for item in IMAGES_FORMAT
        if os.path.splitext(name)[1] == item
    ]

    def image_compose():
        to_image = Image.new(
            "RGB", (IMAGE_COLUMN * IMAGE_SIZE_1, IMAGE_ROW * IMAGE_SIZE_2)
        )

        for y in range(1, IMAGE_ROW + 1):
            for x in range(1, IMAGE_COLUMN + 1):
                from_image = Image.open(
                    IMAGES_PATH + image_names[IMAGE_COLUMN * (y - 1) + x - 1]
                ).resize((IMAGE_SIZE_1, IMAGE_SIZE_2), Image.Resampling.LANCZOS)
                to_image.paste(
                    from_image, ((x - 1) * IMAGE_SIZE_1, (y - 1) * IMAGE_SIZE_2)
                )
        return to_image.save("graph.png")

    image_compose()


def make_tables_img():
    xlsx_path = 'C:/Users/Cloudy/Desktop/MATVEEV/report.xlsx'
    client = win32com.client.Dispatch("Excel.Application")
    wb = client.Workbooks.Open(xlsx_path)

    ws = wb.Worksheets("Статистика по годам")
    ws.Range("A1:E17").CopyPicture(Format=2)
    img = ImageGrab.grabclipboard()
    img.save('to_pdf_1.jpg')

    ws = wb.Worksheets("Статистика по городам")
    ws.Range("A1:E11").CopyPicture(Format=2)
    img = ImageGrab.grabclipboard()
    img.save('to_pdf_2.jpg')

    wb.Close()
    client.Quit()


def create_docx(req):
    docuu = docx.Document()
    style = docuu.styles['Normal']
    style.font.name = 'Verdana'
    style.font.size = Pt(16)
    style.font.bold = True
    docuu.add_paragraph('         Аналитика по зарплатам и городам')
    docuu.add_paragraph(f'                для профессии {req}')
    docuu.add_picture('graph.png', width=docx.shared.Cm(15))
    docuu.add_paragraph('                    Статистика по годам')
    docuu.add_picture('to_pdf_1.jpg', width=docx.shared.Cm(15))
    docuu.add_paragraph()
    docuu.add_paragraph('                  Статистика по городам')
    docuu.add_picture('to_pdf_2.jpg', width=docx.shared.Cm(15))

    docuu.save('generate_pdf.docx')


def full_pdf():
    def file_convert_docx_pdf(p):
        file_in_dir = os.listdir(p)

        for file in file_in_dir:
            if file.endswith('.docx'):
                file_k = f'{file.split(".")[0].replace(".", "_")}.pdf'
                convert(f'{p}\\{file}', f'{p}\\convert_pdf\\{file_k}')
            else:
                continue

    dirs = 'C:/Users/Cloudy/Desktop/MATVEEV/'
    file_convert_docx_pdf(dirs)
    os.startfile(r'C:/Users/Cloudy/Desktop/MATVEEV/convert_pdf/generate_pdf.pdf')


requests = ["Введите название файла: ", "Введите название профессии: ", "Что печатать (Вакансии или Статистика)?: "]
input_inf = [input(input_request) for input_request in requests]
data_set = DataSet(input_inf[0])
if len(data_set.vacancies_full_objects) != 0:
    formatted_inf = formatter_info(data_set.vacancies_full_objects)
    inf = finder_inf(formatted_inf, input_inf[1])
    print_inf(inf, input_inf[1], input_inf[2])

save_all_diagrams()

make_tables_img()

create_docx(input_inf[1])

full_pdf()
